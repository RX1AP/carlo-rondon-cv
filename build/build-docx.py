"""
build-docx.py — Genera 'Carlo Rondon CV.docx' desde data.json
Ejecutar desde la raíz del proyecto: python build/build-docx.py
Requiere: pip install python-docx
"""

import json
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent.parent
DATA_PATH = ROOT / "data.json"
OUT_PATH = ROOT / "Carlo Rondon CV.docx"

# ─── COLORES ───────────────────────────────────────────────────
C_BLACK    = RGBColor(0x0a, 0x0a, 0x0a)
C_DARK     = RGBColor(0x1a, 0x1a, 0x1a)
C_ACCENT   = RGBColor(0x33, 0x33, 0x55)
C_MUTED    = RGBColor(0x55, 0x55, 0x66)
C_BODY     = RGBColor(0x22, 0x22, 0x33)
C_GREEN    = RGBColor(0x16, 0xa3, 0x4a)
C_GOLD     = RGBColor(0xb4, 0x5e, 0x09)


def load():
    with open(DATA_PATH, encoding="utf-8") as f:
        return json.load(f)


def strip_bold(text):
    """Devuelve texto plano sin marcadores **...**."""
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text)


def add_bold_run(para, text):
    """Añade texto con **negrita** como múltiples runs en el párrafo."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
            run.font.color.rgb = C_BLACK
        else:
            run = para.add_run(part)
            run.bold = False
            run.font.color.rgb = C_BODY


def set_font(run, size=10, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=22, bold=True, color=C_BLACK)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    set_font(run, size=8, bold=True, color=C_MUTED)
    # Borde inferior
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), 'DDDDDD')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def h3(doc, text, sub=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=11, bold=True, color=C_BLACK)
    if sub:
        p.add_run("  ")
        r2 = p.add_run(sub)
        set_font(r2, size=9, color=C_MUTED, italic=True)
    return p


def body(doc, text, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    add_bold_run(p, text)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(size)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.5)
    add_bold_run(p, text)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(10)
    return p


def metrics_table(doc, metrics):
    if not metrics:
        return
    table = doc.add_table(rows=1, cols=len(metrics))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for i, m in enumerate(metrics):
        cell = table.rows[0].cells[i]
        cell._tc.get_or_add_tcPr()
        # Valor
        vp = cell.paragraphs[0]
        vp.paragraph_format.space_after = Pt(0)
        vrun = vp.add_run(m["valor"])
        set_font(vrun, size=11, bold=True, color=C_BLACK)
        # Label
        lp = cell.add_paragraph(m["label"])
        lrun = lp.runs[0]
        set_font(lrun, size=8, color=C_MUTED)
        lp.paragraph_format.space_before = Pt(1)
        lp.paragraph_format.space_after = Pt(0)
    # Quitar bordes visibles
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'), 'none')
                tcBorders.append(el)
            tcPr.append(tcBorders)
            # Fondo
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'F5F5F8')
            tcPr.append(shd)
    doc.add_paragraph()


def build(d):
    doc = Document()

    # ─── Márgenes ───────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    p = d["persona"]

    # ─── HEADER ─────────────────────────────────────────────────
    name_p = doc.add_paragraph()
    name_p.paragraph_format.space_after = Pt(2)
    nr = name_p.add_run(f"{p['nombre']} {p['apellido']}")
    set_font(nr, size=26, bold=True, color=C_BLACK)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(6)
    tr = title_p.add_run(p["titulo"])
    set_font(tr, size=12, color=C_MUTED, italic=True)

    contact_p = doc.add_paragraph()
    contact_p.paragraph_format.space_after = Pt(2)
    cr = contact_p.add_run(
        f"{p['ubicacion']}  ·  {p['email']}  ·  {p['whatsapp']}  ·  linkedin.com/in/carlo-rondon"
    )
    set_font(cr, size=9, color=C_MUTED)

    # ─── PERFIL ─────────────────────────────────────────────────
    h2(doc, "Perfil")
    body(doc, d["perfil"])

    # ─── EXPERIENCIA ─────────────────────────────────────────────
    h2(doc, "Experiencia Profesional")

    for exp in d["experiencia"]:
        # Empresa
        ep = doc.add_paragraph()
        ep.paragraph_format.space_before = Pt(10)
        ep.paragraph_format.space_after  = Pt(2)
        er = ep.add_run(exp["empresa"])
        set_font(er, size=12, bold=True, color=C_BLACK)
        pr = ep.add_run(f"   {exp['periodo']}  ·  {exp['duracion']}")
        set_font(pr, size=9, color=C_MUTED)

        for role in exp["roles"]:
            h3(doc, role["titulo"], role.get("tag") or role["periodo"])
            if role.get("tag"):
                pp = doc.add_paragraph()
                pp.paragraph_format.space_before = Pt(0)
                pp.paragraph_format.space_after  = Pt(4)
                pr2 = pp.add_run(role["periodo"])
                set_font(pr2, size=9, color=C_MUTED)

            if role["metricas"]:
                metrics_table(doc, role["metricas"])

            for logro in role["logros"]:
                bullet(doc, logro)

        if exp["skills"]:
            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(6)
            sp.paragraph_format.space_after  = Pt(2)
            sr = sp.add_run("Skills: ")
            set_font(sr, size=9, bold=True, color=C_MUTED)
            sr2 = sp.add_run("  ·  ".join(exp["skills"]))
            set_font(sr2, size=9, color=C_MUTED)

    # ─── EDUCACIÓN ───────────────────────────────────────────────
    h2(doc, "Educación")
    for edu in d["educacion"]:
        ep = doc.add_paragraph()
        ep.paragraph_format.space_before = Pt(6)
        ep.paragraph_format.space_after  = Pt(1)
        er = ep.add_run(edu["titulo"])
        set_font(er, size=11, bold=True, color=C_BLACK)
        ip = doc.add_paragraph()
        ip.paragraph_format.space_before = Pt(0)
        ip.paragraph_format.space_after  = Pt(2)
        ir = ip.add_run(f"{edu['institucion']}  ·  {edu['nivel']}")
        set_font(ir, size=9, color=C_MUTED)

    # ─── CERTIFICACIONES ─────────────────────────────────────────
    h2(doc, "Certificaciones")
    for cert in d["certificaciones"]:
        estado = "En curso" if cert["estado"] == "en_curso" else "✓"
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(3)
        cp.paragraph_format.space_after  = Pt(1)
        cr2 = cp.add_run(f"[{estado}]  ")
        color = C_GOLD if cert["estado"] == "en_curso" else C_GREEN
        set_font(cr2, size=9, bold=True, color=color)
        nr2 = cp.add_run(cert["nombre"])
        set_font(nr2, size=10, bold=True, color=C_BLACK)
        er2 = cp.add_run(f"  —  {cert['entidad']} · {cert['año']}")
        set_font(er2, size=9, color=C_MUTED)

    # ─── IDIOMAS ─────────────────────────────────────────────────
    h2(doc, "Idiomas")
    for lang in d["idiomas"]:
        lp = doc.add_paragraph()
        lp.paragraph_format.space_before = Pt(4)
        lp.paragraph_format.space_after  = Pt(1)
        lr = lp.add_run(f"{lang['nombre']}  —  {lang['nivel']}")
        set_font(lr, size=10, bold=True, color=C_BLACK)
        if lang.get("nota"):
            lr2 = lp.add_run(f"  ({lang['nota']})")
            set_font(lr2, size=9, color=C_MUTED, italic=True)

    # ─── GUARDAR ─────────────────────────────────────────────────
    doc.save(OUT_PATH)
    print(f"[OK] Carlo Rondon CV.docx generado ({OUT_PATH.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    build(load())
